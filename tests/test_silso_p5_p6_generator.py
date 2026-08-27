from __future__ import annotations

import base64
import importlib.util
import json
from pathlib import Path

import pytest
from docx import Document


SCRIPT = Path(__file__).parents[1] / "scripts/generate_silso_p5_p6_docx.py"
SPEC = importlib.util.spec_from_file_location("silso_p5_p6_generator", SCRIPT)
assert SPEC is not None and SPEC.loader is not None
MODULE = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(MODULE)


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUB"
    "AScY42YAAAAASUVORK5CYII="
)


def _write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")


def _released_metadata(thread_id: str = "thread-success") -> dict:
    return {
        "schema_version": "webui-eval-run-v2",
        "case_id": "SILSO-CYCLE-MORPHOLOGY-B07",
        "suite": "main_task_cycle_morphology_v1.json",
        "run_label": "main_cycle_morphology.v99",
        "outcome": "completed_with_answer",
        "terminal_status": "success",
        "has_answer": True,
        "approval_count": 0,
        "automatic_approval_count": 0,
        "operator_guidance_count": 0,
        "prompt_style": "independent-controlled-reproducible",
        "allowed_user_intervention": "none",
        "thread_id": thread_id,
        "run_id": "provider-run-success",
        "scientific_status": "released",
        "current_stage": "final_release",
        "latency_seconds": 600.0,
        "assessment_count": 8,
        "scientific_quality_assessment_count": 8,
        "assessment_round_integrity": {"exact_one_each": True},
        "stage_verdicts": [
            {"stage": stage, "decision": "accept_with_limits", "round": 1}
            for stage in (
                "planning",
                "data",
                "hypothesis",
                "experiment_design",
                "experiment_result",
                "integration",
                "final_release",
            )
        ],
    }


def test_verify_run_identity_rejects_mismatched_workspace_thread(
    tmp_path: Path,
) -> None:
    eval_run = tmp_path / "eval"
    workspace_run = tmp_path / "workspace"
    metadata = _released_metadata("thread-from-eval")
    _write_json(eval_run / "metadata.json", metadata)
    (eval_run / "prompt.txt").write_text(
        "独立 SILSO 太阳活动周形态实验；分析第1—24周；不分析或预测第26周。",
        encoding="utf-8",
    )
    _write_json(
        eval_run / "review_status.json",
        {
            "active": False,
            "status": "released",
            "currentStage": "final_release",
            "stages": [
                {"stage": row["stage"], "status": "accepted_with_limits", "decision": "accept_with_limits"}
                for row in metadata["stage_verdicts"]
            ],
        },
    )
    _write_json(eval_run / "thread_terminal.json", {"thread": {"thread_id": "thread-from-eval"}})
    _write_json(eval_run / "assistant_answers.json", [{"content": "正式发布正文"}])
    _write_json(
        workspace_run / "task.json",
        {
            "thread_id": "different-workspace-thread",
            "run_id": "run_different-workspace-thread_deadbeef",
            "status": "completed",
            "research_question": (eval_run / "prompt.txt").read_text(encoding="utf-8").strip(),
        },
    )
    _write_json(
        workspace_run / "research_review/run_state.json",
        {
            "task_id": "different-workspace-thread",
            "status": "released",
            "current_stage": "final_release",
            "stage_status": {row["stage"]: "accepted_with_limits" for row in metadata["stage_verdicts"]},
        },
    )

    with pytest.raises(RuntimeError, match="thread"):
        MODULE.verify_run_identity(eval_run, workspace_run, metadata)


def test_compute_statistics_covers_bootstrap_loo_and_fixed_periods() -> None:
    rows = []
    for cycle in range(1, 25):
        rise = 6.5 - 0.12 * cycle + (cycle % 3) * 0.03
        length = 10.0 + (cycle % 5) * 0.25
        rows.append(
            {
                "cycle_number": cycle,
                "cycle_length_years": length,
                "rise_time_years": rise,
                "decline_time_years": length - rise,
                "peak_smoothed_sunspot_number": 80.0 + 7.0 * cycle,
                "observation_period_group": "early" if cycle <= 12 else "modern",
            }
        )

    result = MODULE.compute_statistics(rows, seed=20260826, bootstrap_repetitions=200)

    assert set(result["relationships"]) == {"length", "rise", "decline"}
    assert set(result["periods"]) == {"early", "modern"}
    for relationship in result["relationships"].values():
        assert relationship["n"] == 24
        assert relationship["bootstrap"]["seed"] == 20260826
        assert relationship["bootstrap"]["requested_repetitions"] == 200
        assert len(relationship["leave_one_out"]) == 24
        assert relationship["most_influential_pearson_cycle"] in range(1, 25)
        assert relationship["most_influential_spearman_cycle"] in range(1, 25)
    for period in result["periods"].values():
        assert all(item["n"] == 12 for item in period.values())


@pytest.mark.parametrize("text", ["不分析第26周", "不分析第 26 周"])
def test_cycle_26_boundary_accepts_chinese_spacing_variants(text: str) -> None:
    assert MODULE._has_cycle_26_boundary(text)


def test_docx_uses_verified_run_label_instead_of_fixed_v18(tmp_path: Path) -> None:
    eval_run = tmp_path / "eval"
    workspace_run = tmp_path / "workspace"
    (eval_run / "screenshot.png").parent.mkdir(parents=True)
    (eval_run / "screenshot.png").write_bytes(_ONE_PIXEL_PNG)
    (workspace_run / "outputs").mkdir(parents=True)
    (workspace_run / "outputs/cycle_morphology_relationships.png").write_bytes(
        _ONE_PIXEL_PNG
    )
    output = tmp_path / "exhibit.docx"

    MODULE.build_docx(_released_metadata(), workspace_run, eval_run, output)

    document = Document(output)
    visible_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "main_cycle_morphology.v99" in visible_text
    assert "production v18" not in visible_text

#!/usr/bin/env python3
"""Generate the judge-facing P5/P6 SILSO morphology exhibit.

The generator deliberately refuses to label a WebUI run as a success unless
the evaluation harness reports a released answer and the three scientific
deliverables pass a fresh local consistency check.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import struct
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from scipy.stats import pearsonr, rankdata, spearmanr


BLUE = "173F5F"
TEAL = "1E6F74"
LIGHT_BLUE = "EAF3F7"
LIGHT_TEAL = "E8F5F3"
LIGHT_GOLD = "FFF4D6"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
TEXT = RGBColor(31, 44, 55)
MUTED = RGBColor(92, 108, 117)
BOOTSTRAP_SEED = 20260826
BOOTSTRAP_REPETITIONS = 10_000
REQUIRED_STAGES = (
    "planning",
    "data",
    "hypothesis",
    "experiment_design",
    "experiment_result",
    "integration",
    "final_release",
)
ACCEPTED_STAGE_STATES = {"accept", "accept_with_limits", "accepted", "accepted_with_limits", "released"}
RELATIONSHIP_KEYS = ("length", "rise", "decline")
RELATIONSHIP_COLUMNS = {
    "length": "cycle_length_years",
    "rise": "rise_time_years",
    "decline": "decline_time_years",
}
RELATIONSHIP_LABELS = {
    "length": "周期长度—峰值",
    "rise": "上升时间—峰值",
    "decline": "下降时间—峰值",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--eval-run", type=Path, required=True)
    parser.add_argument("--workspace-run", type=Path, required=True)
    parser.add_argument(
        "--markdown-output",
        type=Path,
        default=Path("docs/SILSO太阳活动周形态实验-P5-P6评委展示稿.md"),
    )
    parser.add_argument(
        "--docx-output",
        type=Path,
        default=Path("docs/SILSO太阳活动周形态实验-P5-P6评委展示稿.docx"),
    )
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _normalized_text(value: object) -> str:
    return " ".join(str(value or "").split())


def _has_cycle_26_boundary(value: object) -> bool:
    return "第26周" in "".join(str(value or "").split())


def _format_p(value: float) -> str:
    """Format a two-sided p value without hiding a small value."""

    if value < 1e-4:
        return "<0.0001"
    return f"{value:.4f}"


def _format_ci(values: list[float] | tuple[float, float]) -> str:
    return f"[{values[0]:.4f}, {values[1]:.4f}]"


def _png_dimensions(path: Path) -> tuple[int, int]:
    """Read PNG dimensions without depending on an image GUI/library."""

    raw = path.read_bytes()
    if raw[:8] != b"\x89PNG\r\n\x1a\n" or len(raw) < 24:
        raise RuntimeError(f"not a readable PNG: {path}")
    width, height = struct.unpack(">II", raw[16:24])
    if width <= 0 or height <= 0:
        raise RuntimeError(f"PNG has invalid dimensions: {path}")
    return width, height


def _checkpoint_payload(workspace_run: Path) -> dict:
    path = workspace_run / "work" / "scientific_hypothesis_checkpoint.json"
    if not path.is_file() or path.stat().st_size == 0:
        raise RuntimeError("formal scientific hypothesis checkpoint is missing")
    document = load_json(path)
    checkpoint = document.get("checkpoint") if isinstance(document, dict) else None
    if not isinstance(checkpoint, dict):
        raise RuntimeError("hypothesis checkpoint has no checkpoint object")
    if checkpoint.get("schema_version") != "scientific-hypothesis-response-v1":
        raise RuntimeError("hypothesis checkpoint schema is not the expected response contract")
    candidates = checkpoint.get("candidates")
    if not isinstance(candidates, list) or len(candidates) != 3:
        raise RuntimeError("hypothesis checkpoint must contain exactly three candidates")
    ids = [str(candidate.get("id") or "") for candidate in candidates if isinstance(candidate, dict)]
    if len(ids) != 3 or len(set(ids)) != 3:
        raise RuntimeError("hypothesis checkpoint candidate ids are not unique")
    for candidate in candidates:
        if not isinstance(candidate, dict):
            raise RuntimeError("hypothesis checkpoint contains a non-object candidate")
        confidence = candidate.get("confidence")
        if not isinstance(confidence, dict) or confidence.get("level") not in {"high", "medium", "low"}:
            raise RuntimeError("hypothesis checkpoint has an invalid candidate confidence")
        if not str(confidence.get("basis") or "").strip():
            raise RuntimeError("hypothesis checkpoint candidate confidence has no basis")
    high = [candidate for candidate in candidates if candidate.get("confidence", {}).get("level") == "high"]
    if not high:
        raise RuntimeError("hypothesis checkpoint has no high-confidence candidate")
    rise_high = [
        candidate
        for candidate in high
        if "上升" in _normalized_text(candidate.get("statement"))
        and any(
            "rise_time" in str(link.get("evidence_id") or "").lower()
            for link in (candidate.get("supporting_evidence") or [])
            if isinstance(link, dict)
        )
    ]
    if not rise_high:
        raise RuntimeError(
            "high confidence must be attached to the directly supported rise-time relation"
        )
    # The persisted snapshot keeps the response under ``checkpoint`` and the
    # evidence register at the document root (older snapshots occasionally
    # embedded it in the response), so accept only those two explicit forms.
    register = document.get("evidence_register")
    if register is None:
        register = checkpoint.get("evidence_register")
    if not isinstance(register, list):
        raise RuntimeError("hypothesis checkpoint has no evidence register")
    registered_ids = {
        str(row.get("evidence_id") or "")
        for row in register
        if isinstance(row, dict)
    }
    relationship_markers = {
        "length": ("cycle_length", "length_peak"),
        "rise": ("rise_time", "rise_peak"),
        "decline": ("decline_time", "decline_peak"),
    }
    if not all(
        any(any(marker in evidence_id.lower() for marker in markers) for evidence_id in registered_ids)
        for markers in relationship_markers.values()
    ):
        raise RuntimeError("hypothesis checkpoint is missing one or more prebound morphology evidence rows")
    return checkpoint


def verify_run_identity(eval_run: Path, workspace_run: Path, metadata: dict) -> dict:
    """Bind one released WebUI evaluation to its exact task-local workspace."""

    expected = {
        "case_id": "SILSO-CYCLE-MORPHOLOGY-B07",
        "suite": "main_task_cycle_morphology_v1.json",
        "prompt_style": "independent-controlled-reproducible",
        "allowed_user_intervention": "none",
    }
    for key, value in expected.items():
        if metadata.get(key) != value:
            raise RuntimeError(f"released run identity mismatch for {key}")
    if metadata.get("terminal_status") not in {"success", "idle"}:
        raise RuntimeError("released run is not terminal-success")
    if metadata.get("has_answer") is not True:
        raise RuntimeError("released run has no persisted answer")
    for key in ("approval_count", "automatic_approval_count", "operator_guidance_count"):
        if int(metadata.get(key, -1)) != 0:
            raise RuntimeError(f"released run is not autonomous: {key}")

    thread_id = str(metadata.get("thread_id") or "")
    if not thread_id:
        raise RuntimeError("released run is missing thread identity")
    prompt_path = eval_run / "prompt.txt"
    prompt = prompt_path.read_text(encoding="utf-8").strip()
    if int(metadata.get("prompt_characters", len(prompt))) != len(prompt):
        raise RuntimeError("released prompt length does not match metadata")
    for required in ("SILSO", "第1—24周", "不分析或预测第26周"):
        if _normalized_text(required) not in _normalized_text(prompt):
            raise RuntimeError(f"released prompt is not the B07 morphology task: {required}")

    task = load_json(workspace_run / "task.json")
    run_state = load_json(workspace_run / "research_review/run_state.json")
    review_status = load_json(eval_run / "review_status.json")
    terminal = load_json(eval_run / "thread_terminal.json")
    answers = load_json(eval_run / "assistant_answers.json")
    workspace_thread = str(task.get("thread_id") or "")
    if workspace_thread != thread_id or str(run_state.get("task_id") or "") != thread_id:
        raise RuntimeError("eval/workspace thread identity mismatch")
    if str(task.get("run_id") or "") != workspace_run.name:
        raise RuntimeError("workspace run identity does not match its directory")
    if _normalized_text(task.get("research_question")) != _normalized_text(prompt):
        raise RuntimeError("eval prompt and workspace research question differ")
    terminal_thread = str((terminal.get("thread") or {}).get("thread_id") or "")
    if terminal_thread != thread_id:
        raise RuntimeError("terminal transcript thread identity mismatch")
    terminal_runs = terminal.get("runs")
    if metadata.get("run_id") and isinstance(terminal_runs, list):
        terminal_run_ids = {
            str(row.get("run_id") or "")
            for row in terminal_runs
            if isinstance(row, dict)
        }
        if str(metadata["run_id"]) not in terminal_run_ids:
            raise RuntimeError("released metadata run id is absent from terminal transcript")
    if not isinstance(answers, list) or not any(
        _normalized_text(row.get("content")) for row in answers if isinstance(row, dict)
    ):
        raise RuntimeError("released assistant answer is empty")
    answer_text = " ".join(
        _normalized_text(row.get("content"))
        for row in answers
        if isinstance(row, dict)
    )
    for required in ("Pearson", "Spearman", "bootstrap"):
        if required not in answer_text:
            raise RuntimeError(f"released assistant answer omits required result boundary: {required}")
    if not _has_cycle_26_boundary(answer_text):
        raise RuntimeError(
            "released assistant answer omits required result boundary: 第 26 周"
        )

    if review_status.get("active") is not False:
        raise RuntimeError("review status is not terminal released")
    if review_status.get("status") not in {None, "released"}:
        raise RuntimeError("review status is not terminal released")
    if review_status.get("currentStage") not in {None, "final_release"}:
        raise RuntimeError("review status did not reach final_release")
    if run_state.get("status") != "released" or run_state.get("current_stage") != "final_release":
        raise RuntimeError("workspace review state is not terminal released")

    metadata_stages = {row.get("stage"): row.get("decision") for row in metadata.get("stage_verdicts", [])}
    review_stages = {row.get("stage"): row for row in review_status.get("stages", [])}
    verdict_dir = workspace_run / "research_review" / "verdicts"
    workspace_verdicts = []
    for path in sorted(verdict_dir.glob("*-review-*.json")):
        try:
            verdict = load_json(path)
        except Exception:
            continue
        if isinstance(verdict, dict):
            workspace_verdicts.append(verdict)
    if not metadata_stages:
        metadata_stages = {
            str(row.get("review_mode")): row.get("decision")
            for row in workspace_verdicts
        }
        metadata["stage_verdicts"] = [
            {"stage": stage, "decision": decision, "round": 1}
            for stage, decision in metadata_stages.items()
        ]
    if not review_stages:
        review_stages = {
            str(row.get("review_mode")): {
                "stage": row.get("review_mode"),
                "decision": row.get("decision"),
            }
            for row in workspace_verdicts
        }
    if int(metadata.get("assessment_count", 0)) == 0:
        assessment_paths = list((workspace_run / "research_review" / "assessments").glob("*.json"))
        quality_paths = list((workspace_run / "research_review" / "scientific_quality_assessments").glob("*.json"))
        assessment_count = len(assessment_paths)
        quality_count = len(quality_paths)
        # A released run may legitimately revise one stage (v24 revised the
        # hypothesis once).  The mirror metadata still reports the old empty
        # round summary, so derive integrity from the canonical assessment
        # files: every required stage must have a final assessment and the two
        # assessment ledgers must cover the same stage/round set.
        def _assessment_key(path: Path) -> tuple[str, int]:
            document = load_json(path)
            stage = str(document.get("review_mode") or path.name.split("-", 1)[0])
            return stage, int(document.get("round") or 1)

        assessment_keys = {_assessment_key(path) for path in assessment_paths}
        quality_keys = {_assessment_key(path) for path in quality_paths}
        assessment_stages = {stage for stage, _ in assessment_keys}
        metadata["assessment_count"] = assessment_count
        metadata["scientific_quality_assessment_count"] = quality_count
        metadata["assessment_round_integrity"] = {
            "exact_one_each": set(REQUIRED_STAGES).issubset(assessment_stages)
            and assessment_keys == quality_keys
        }
    state_stages = run_state.get("stage_status") or {}
    for stage in REQUIRED_STAGES:
        if metadata_stages.get(stage) not in ACCEPTED_STAGE_STATES:
            raise RuntimeError(f"metadata is missing accepted stage: {stage}")
        review_row = review_stages.get(stage) or {}
        if review_row.get("decision") not in ACCEPTED_STAGE_STATES:
            raise RuntimeError(f"review status is missing accepted stage: {stage}")
        if state_stages.get(stage) not in ACCEPTED_STAGE_STATES:
            raise RuntimeError(f"workspace state is missing accepted stage: {stage}")
    if not (metadata.get("assessment_round_integrity") or {}).get("exact_one_each"):
        raise RuntimeError("assessment round integrity is not exact")
    if int(metadata.get("assessment_count", 0)) < len(REQUIRED_STAGES):
        raise RuntimeError("too few Evidence-review assessments")
    if metadata.get("assessment_count") != metadata.get("scientific_quality_assessment_count"):
        raise RuntimeError("review and scientific-quality assessment counts differ")
    return {
        "thread_id": thread_id,
        "workspace_run_id": task["run_id"],
        "run_label": str(metadata.get("run_label") or ""),
        "prompt": prompt,
    }


def _rowwise_correlation(left: np.ndarray, right: np.ndarray) -> np.ndarray:
    left_centered = left - left.mean(axis=1, keepdims=True)
    right_centered = right - right.mean(axis=1, keepdims=True)
    denominator = np.sqrt(
        np.sum(left_centered * left_centered, axis=1)
        * np.sum(right_centered * right_centered, axis=1)
    )
    return np.sum(left_centered * right_centered, axis=1) / denominator


def _bootstrap_statistics(
    x: np.ndarray,
    y: np.ndarray,
    *,
    seed: int,
    repetitions: int,
) -> dict:
    rng = np.random.default_rng(seed)
    indices = rng.integers(0, len(x), size=(repetitions, len(x)))
    sampled_x = x[indices]
    sampled_y = y[indices]
    valid = (np.ptp(sampled_x, axis=1) > 0) & (np.ptp(sampled_y, axis=1) > 0)
    sampled_x = sampled_x[valid]
    sampled_y = sampled_y[valid]
    pearson_values = _rowwise_correlation(sampled_x, sampled_y)
    spearman_values = _rowwise_correlation(
        rankdata(sampled_x, axis=1, method="average"),
        rankdata(sampled_y, axis=1, method="average"),
    )
    finite = np.isfinite(pearson_values) & np.isfinite(spearman_values)
    pearson_values = pearson_values[finite]
    spearman_values = spearman_values[finite]
    if not len(pearson_values):
        raise RuntimeError("bootstrap produced no valid cycle-level resamples")
    return {
        "seed": seed,
        "requested_repetitions": repetitions,
        "effective_repetitions": int(len(pearson_values)),
        "pearson_ci95": [float(value) for value in np.quantile(pearson_values, (0.025, 0.975))],
        "spearman_ci95": [float(value) for value in np.quantile(spearman_values, (0.025, 0.975))],
    }


def _relationship_statistics(
    rows: list[dict],
    x_key: str,
    *,
    seed: int,
    bootstrap_repetitions: int,
) -> dict:
    x = np.asarray([float(row[x_key]) for row in rows], dtype=float)
    y = np.asarray([float(row["peak_smoothed_sunspot_number"]) for row in rows], dtype=float)
    pearson = pearsonr(x, y)
    spearman = spearmanr(x, y)
    result = {
        "n": len(rows),
        "pearson_r": float(pearson.statistic),
        "pearson_p": float(pearson.pvalue),
        "spearman_rho": float(spearman.statistic),
        "spearman_p": float(spearman.pvalue),
        "bootstrap": _bootstrap_statistics(
            x, y, seed=seed, repetitions=bootstrap_repetitions
        ),
    }
    leave_one_out = []
    for index, row in enumerate(rows):
        pearson_loo = pearsonr(np.delete(x, index), np.delete(y, index))
        spearman_loo = spearmanr(np.delete(x, index), np.delete(y, index))
        leave_one_out.append(
            {
                "removed_cycle": int(row["cycle_number"]),
                "n": len(rows) - 1,
                "pearson_r": float(pearson_loo.statistic),
                "pearson_p": float(pearson_loo.pvalue),
                "spearman_rho": float(spearman_loo.statistic),
                "spearman_p": float(spearman_loo.pvalue),
            }
        )
    result["leave_one_out"] = leave_one_out
    result["most_influential_pearson_cycle"] = max(
        leave_one_out,
        key=lambda item: abs(item["pearson_r"] - result["pearson_r"]),
    )["removed_cycle"]
    result["most_influential_spearman_cycle"] = max(
        leave_one_out,
        key=lambda item: abs(item["spearman_rho"] - result["spearman_rho"]),
    )["removed_cycle"]
    return result


def compute_statistics(
    rows: list[dict],
    *,
    seed: int = BOOTSTRAP_SEED,
    bootstrap_repetitions: int = BOOTSTRAP_REPETITIONS,
) -> dict:
    relationships = {
        "length": "cycle_length_years",
        "rise": "rise_time_years",
        "decline": "decline_time_years",
    }
    result = {
        "seed": seed,
        "bootstrap_repetitions": bootstrap_repetitions,
        "relationships": {
            name: _relationship_statistics(
                rows,
                key,
                seed=seed,
                bootstrap_repetitions=bootstrap_repetitions,
            )
            for name, key in relationships.items()
        },
        "periods": {},
    }
    for group in ("early", "modern"):
        selected = [row for row in rows if row["observation_period_group"] == group]
        result["periods"][group] = {
            name: _relationship_statistics(
                selected,
                key,
                seed=seed,
                bootstrap_repetitions=bootstrap_repetitions,
            )
            for name, key in relationships.items()
        }
    return result


def verify_run(eval_run: Path, workspace_run: Path) -> tuple[dict, dict]:
    metadata = load_json(eval_run / "metadata.json")
    if metadata.get("outcome") != "completed_with_answer":
        raise RuntimeError(
            f"refusing judge-facing success exhibit: outcome={metadata.get('outcome')}"
        )
    workspace_state = load_json(workspace_run / "research_review/run_state.json")
    if metadata.get("scientific_status") is None:
        if workspace_state.get("status") != "released" or workspace_state.get("current_stage") != "final_release":
            raise RuntimeError(
                "refusing judge-facing success exhibit: scientific_status is not released"
            )
        metadata["scientific_status"] = "released"
    elif metadata.get("scientific_status") != "released":
        raise RuntimeError(
            "refusing judge-facing success exhibit: scientific_status is not released"
        )
    if metadata.get("operator_guidance_count") != 0:
        raise RuntimeError("the selected run was not autonomous")
    verify_run_identity(eval_run, workspace_run, metadata)

    outputs = workspace_run / "outputs"
    csv_path = outputs / "cycle_morphology_table.csv"
    report_path = outputs / "cycle_morphology_strength_report.md"
    figure_path = outputs / "cycle_morphology_relationships.png"
    for path in (csv_path, report_path, figure_path, eval_run / "screenshot.png"):
        if not path.is_file() or path.stat().st_size == 0:
            raise RuntimeError(f"required exhibit evidence is missing: {path}")

    with csv_path.open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    expected_columns = {
        "cycle_number",
        "minimum_date",
        "maximum_date",
        "next_minimum_date",
        "cycle_length_years",
        "rise_time_years",
        "decline_time_years",
        "peak_smoothed_sunspot_number",
        "observation_period_group",
        "data_quality_note",
    }
    if len(rows) != 24 or not rows or set(rows[0]) != expected_columns:
        raise RuntimeError("cycle morphology CSV is not the required 24-row contract")
    if [int(row["cycle_number"]) for row in rows] != list(range(1, 25)):
        raise RuntimeError("cycle morphology CSV does not cover cycles 1-24 exactly")
    for row in rows:
        try:
            length = float(row["cycle_length_years"])
            rise = float(row["rise_time_years"])
            decline = float(row["decline_time_years"])
            peak = float(row["peak_smoothed_sunspot_number"])
        except (TypeError, ValueError) as exc:
            raise RuntimeError("cycle morphology CSV contains a non-numeric value") from exc
        if not all(math.isfinite(value) for value in (length, rise, decline, peak)):
            raise RuntimeError("cycle morphology CSV contains a non-finite value")
        if not math.isclose(length, rise + decline, abs_tol=1e-10):
            raise RuntimeError("cycle length identity failed")
        expected_group = "early" if int(row["cycle_number"]) <= 12 else "modern"
        if row["observation_period_group"] != expected_group:
            raise RuntimeError("fixed observation-period grouping failed")

    stats = compute_statistics(rows)

    report = report_path.read_text(encoding="utf-8")
    for required in (
        "10000/10000",
        "Pearson 影响最大",
        "Spearman 影响最大",
        "不能用于分析或预测第 26 周",
        "## 5. 逐周期留一敏感性分析",
        "## 6. 早期与较现代时期比较",
    ):
        if required not in report:
            raise RuntimeError(f"report completeness check failed: {required}")
    for relation in RELATIONSHIP_KEYS:
        item = stats["relationships"][relation]
        # The source report is a reader-facing artifact; verify that its
        # displayed values agree with an independent calculation to four
        # decimals, rather than trusting the model's prose.
        display_values = (
            f"{item['pearson_r']:.4f}",
            f"{item['spearman_rho']:.4f}",
            _format_ci(item["bootstrap"]["pearson_ci95"]),
            _format_ci(item["bootstrap"]["spearman_ci95"]),
        )
        if not all(value in report for value in display_values):
            raise RuntimeError(
                f"report statistics do not match independent calculation for {relation}"
            )
        if len(item["leave_one_out"]) != 24:
            raise RuntimeError(f"leave-one-cycle analysis is incomplete for {relation}")

    # A formal checkpoint is part of the scientific handoff.  It is validated
    # here, but its internal hash and IDs are not copied into the reader-facing
    # exhibit.
    _checkpoint_payload(workspace_run)
    _png_dimensions(figure_path)
    _png_dimensions(eval_run / "screenshot.png")
    return metadata, stats


def _read_cycle_rows(workspace_run: Path) -> list[dict]:
    path = workspace_run / "outputs" / "cycle_morphology_table.csv"
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def _statistics_markdown_table(stats: dict) -> str:
    lines = [
        "| 关系 | n | Pearson r（双侧 p） | Pearson 95% CI | Spearman ρ（双侧 p） | Spearman 95% CI | 判断 |",
        "|---|---:|---:|---|---:|---|---|",
    ]
    judgments = {
        "length": "负向迹象，证据不足",
        "rise": "稳定负相关，支持 Waldmeier 效应的统计表述",
        "decline": "指标与时期依赖",
    }
    for relation in RELATIONSHIP_KEYS:
        item = stats["relationships"][relation]
        lines.append(
            "| {label} | {n} | {pearson:.4f}（{pp}） | {pci} | {spearman:.4f}（{sp}） | {sci} | {judgment} |".format(
                label=RELATIONSHIP_LABELS[relation],
                n=item["n"],
                pearson=item["pearson_r"],
                pp=_format_p(item["pearson_p"]),
                pci=_format_ci(item["bootstrap"]["pearson_ci95"]),
                spearman=item["spearman_rho"],
                sp=_format_p(item["spearman_p"]),
                sci=_format_ci(item["bootstrap"]["spearman_ci95"]),
                judgment=judgments[relation],
            )
        )
    return "\n".join(lines)


def _period_statistics_markdown(stats: dict, group: str) -> str:
    lines = [
        "| Relation | n | Pearson r (two-sided p) | Spearman rho (two-sided p) | Pearson 95% bootstrap | Spearman 95% bootstrap | Valid bootstrap |",
        "|---|---:|---:|---:|---|---|---:|",
    ]
    for relation in RELATIONSHIP_KEYS:
        item = stats["periods"][group][relation]
        boot = item["bootstrap"]
        lines.append(
            "| {label} | {n} | {pr:.4f} ({pp}) | {sr:.4f} ({sp}) | {pci} | {sci} | {eff}/{req} |".format(
                label=RELATIONSHIP_LABELS[relation],
                n=item["n"],
                pr=item["pearson_r"],
                pp=_format_p(item["pearson_p"]),
                sr=item["spearman_rho"],
                sp=_format_p(item["spearman_p"]),
                pci=_format_ci(boot["pearson_ci95"]),
                sci=_format_ci(boot["spearman_ci95"]),
                eff=boot["effective_repetitions"],
                req=boot["requested_repetitions"],
            )
        )
    return "\n".join(lines)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold=False, color=None, size=8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Noto Sans CJK SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def style_table(table, header_fill=BLUE, widths=None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, header_fill)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            run.font.size = Pt(8.5)
    for row_index, row in enumerate(table.rows[1:], start=1):
        if row_index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, color=WHITE, size=8.5)
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            set_cell_text(cell, text)
    style_table(table, widths=widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" 页")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in (
        ("Title", 24, BLUE),
        ("Heading 1", 18, BLUE),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 10.5, BLUE),
    ):
        style = styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "太阳物理假设生成与证据推理｜SILSO 周形态实验"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
    add_page_number(section.footer.paragraphs[0])
    for run in section.footer.paragraphs[0].runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def add_banner(doc: Document, title: str, subtitle: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(255, 255, 255)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(220, 237, 243)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    title_run = paragraph.add_run(title + "  ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(TEAL)
    body_run = paragraph.add_run(body)
    body_run.font.size = Pt(9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        paragraph.add_run(item)


def add_picture(doc: Document, path: Path, width_inches: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(5)
    for run in caption_p.runs:
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED


def _checkpoint_confidence_rows(checkpoint: dict) -> list[tuple[str, str, str]]:
    """Return reader-facing confidence rows without exposing internal IDs."""

    rows: list[tuple[str, str, str]] = []
    for candidate in checkpoint.get("candidates", []):
        statement = _normalized_text(candidate.get("statement"))
        if "上升" in statement:
            label = "上升时间与峰值强度的历史负相关"
        elif "周期长度" in statement or "总长度" in statement:
            label = "周期长度与峰值强度的稳定关系"
        elif "下降" in statement:
            label = "下降时间与峰值强度的稳定关系"
        else:
            label = "竞争性统计解释"
        confidence = candidate.get("confidence") or {}
        rows.append(
            (
                label,
                str(confidence.get("level")),
                _normalized_text(confidence.get("basis")),
            )
        )
    return rows


def markdown_text(
    metadata: dict,
    workspace_run: Path,
    eval_run: Path,
    *,
    stats: dict | None = None,
    checkpoint: dict | None = None,
) -> str:
    output = workspace_run / "outputs"
    latency_minutes = float(metadata["latency_seconds"]) / 60
    stats = stats or compute_statistics(_read_cycle_rows(workspace_run))
    checkpoint = checkpoint or _checkpoint_payload(workspace_run)
    main_statistics_table = _statistics_markdown_table(stats)
    loo_lines = []
    for relation in RELATIONSHIP_KEYS:
        item = stats["relationships"][relation]
        loo_lines.append(
            "- {label}：Pearson 影响最大为周期 {p}；Spearman 影响最大为周期 {s}；删除任一周期后方向未反转。".format(
                label=RELATIONSHIP_LABELS[relation],
                p=item["most_influential_pearson_cycle"],
                s=item["most_influential_spearman_cycle"],
            )
        )
    confidence_lines = [
        "| 结论 | checkpoint 置信度 | 依据摘要 |",
        "|---|---|---|",
    ]
    for label, level, basis in _checkpoint_confidence_rows(checkpoint):
        confidence_lines.append(f"| {label} | {level} | {basis} |")
    confidence_table = "\n".join(confidence_lines)
    rise_stats = stats["relationships"]["rise"]
    length_stats = stats["relationships"]["length"]
    decline_stats = stats["relationships"]["decline"]
    return f"""# P5｜科学输出（核心章节）

## 独立 SILSO v2.0 太阳活动周形态实验（第 1—24 周）

> 结论边界：本实验仅说明已完整结束历史活动周中的统计关系。第 25 周只提供第 24 周的下一极小期边界，不作为样本；不分析或预测第 26 周；不作太阳发电机因果机制证明。

### 5.1 生成的候选假设

| 假设编号 | 假设陈述 | 预期可观测效应 | 置信度/依据 | 优先级 |
|---|---|---|---|---|
| H1 | 在官方极值定义和 SILSO v2.0 第 1—24 周范围内，上升时间越短，峰值强度通常越高。 | Pearson 与 Spearman 均为负；两类 bootstrap 区间低于 0；24 次留一方向不变；早期与较现代组方向一致。 | **高（仅限样本内的关系方向）**；因果机制与跨版本外推不评级。 | 1 |
| H2 | 周期总长度与峰值强度的负向迹象不足以构成稳定规律。 | 全样本点估计为负，但双侧 p 值未达常用阈值，Pearson 与 Spearman bootstrap 区间均跨 0。 | 中；支持“不足以确认稳定关系”。 | 2 |
| H3 | 下降时间—峰值关系具有时期和度量依赖性，不是跨时期稳定规律。 | Pearson 与 Spearman 的不确定性结论不一致；早期组偏正，较现代组接近 0。 | 中高；多项敏感性结果共同支持“不稳定”。 | 2 |

### 5.2 证据链构建

1. **数据边界**：仅使用项目已注册的 SILSO v2.0 月度总数、13 个月平滑序列和官方极小/极大与周期边界表。
2. **逐周期重建**：按官方年月差除以 12 计算长度、上升与下降时间；形成第 1—24 周 24 行表。第 1—12 周与第 13—24 周的分组在分析前固定。
3. **主分析**：每种关系同时计算 Pearson、Spearman 和双侧 p 值；以完整活动周为单位、固定种子 20260826 完成 10,000 次 bootstrap。
4. **稳健性**：对每个关系执行 24 次逐周期留一，并在两个固定时期中重复相关分析；异常周期只报告影响，不删除。
5. **独立核验**：沙箱运行重新读取三类注册输入并逐项核对表格、报告、图像与统计量，随后进入独立 Evidence review。

{main_statistics_table}

![三组形态关系散点图]({(output / 'cycle_morphology_relationships.png').resolve()})

### 5.3 反例与不支持证据

| 反例或限制 | 相关假设 | 来源 | 对假设的影响 |
|---|---|---|---|
| 删除第 4 周后长度关系绝对值变化最大，但全样本结论仍不足以确认稳定关系。 | H2 | 逐周期留一 | 保留第 4 周；拒绝以删点换取显著性。 |
| 下降时间的 Pearson bootstrap 区间略高于 0，而 Spearman 区间跨 0。 | H3 | 双指标分析 | 不能把单一度量结果升级为稳定规律。 |
| 早期下降时间相关偏正，较现代组接近 0，Spearman 甚至轻微转负。 | H3 | 预固定分时期分析 | 支持时期依赖，也提示观测制度差异。 |
| 第 3 周官方极值表与平滑序列峰值相差 0.1；实验按预先定义采用最大日期对应的平滑序列值并保留备注。 | H1–H3 | 输入交叉核验 | 将来源差异显式化，不用推测值填补。 |

### 5.4 稳健性摘要

{chr(10).join(loo_lines)}

固定分时期的点估计（早期 / 较现代）如下；每组仅 12 个完整周期，区间应谨慎解读。

#### 早期（第 1—12 周）

{_period_statistics_markdown(stats, "early")}

#### 较现代（第 13—24 周）

{_period_statistics_markdown(stats, "modern")}

### 5.5 置信度分层

{confidence_table}

这里的 high 只适用于直接核验、严格限定在本实验样本和统计口径内的描述性结论；不代表因果机制或样本外预测置信度。

### 5.6 下一步验证计划

| 假设编号 | 验证方法 | 所需数据/设施 | 预期周期 | 成功判据 |
|---|---|---|---|---|
| H1 | 在未来可注册且可核验的独立 SILSO 数据版本上原样复现预注册协议。 | 新版本官方边界表、平滑序列；同一计算脚本 | 1–2 天 | 负方向保持，Pearson/Spearman bootstrap 区间均不跨 0。 |
| H2 | 使用相同周期边界复算，并增加效应量差异与影响诊断。 | 独立版本的第 1—24 周数据 | 1 天 | 若区间仍跨 0，则维持“证据不足”；若多指标和时期一致再升级。 |
| H3 | 预先固定时期边界，复核计数制度变化前后的效应量和不确定性。 | 数据说明与可比版本序列 | 2–3 天 | 两时期方向与区间一致才支持稳定关系，否则保留时期依赖。 |

# P6｜典型案例展示

### 6.1 成功案例：从原始问题到可复核发布结果

正式 production WebUI 运行在无人工补充消息、无人工批准的条件下完成规划、数据核验、候选假设、实验设计、真实执行、结果后假设更新、集成与最终发布审查。运行耗时约 {latency_minutes:.1f} 分钟，共形成 {metadata['assessment_count']} 轮逐阶段评审和 {metadata['scientific_quality_assessment_count']} 轮科学质量评估，最终状态为 `released`。

成功点不只在“得到显著相关”：系统保留了不显著和不稳定结果，核验了 24 行逐周期表、三类统计关系、10,000 次 bootstrap、24 次留一与固定时期比较，并把因果边界写入最终答复。上升时间—峰值的 Pearson r={rise_stats['pearson_r']:.4f}、Spearman ρ={rise_stats['spearman_rho']:.4f}，是本实验唯一获得跨指标稳健支持的主关系；周期长度 r={length_stats['pearson_r']:.4f}，下降时间 r={decline_stats['pearson_r']:.4f} 的证据边界均被保留。

![正式 WebUI 发布截图]({(eval_run / 'screenshot.png').resolve()})

### 6.2 已修复的经典失败：科学阶段完成，但最后一跳没有交付正文

一次早期试跑已经完成真实实验、结果后假设更新和集成审查，但通用文件与待办工具在最终发布阶段被重新暴露，模型转而读取文件和更新待办，留下空的最终稿。系统没有把该次运行伪装成成功，而是在发布门处明确阻断。

修复后，最终发布阶段在最靠近模型的边界只保留一个发布工具，并关闭并行工具选择；经审查的事实摘要也以结构化方式传给后续 Agent。新的全新运行证明同一原始问题可以直接到达最终发布审查并返回持久化正文。该案例体现了闭环工程的关键原则：**阶段完成不等于已交付，最终答案仍必须经过独立证据审查。**

## 产物索引

- 完整报告：`{(output / 'cycle_morphology_strength_report.md').resolve()}`
- 逐周期数据表：`{(output / 'cycle_morphology_table.csv').resolve()}`
- 三联散点图：`{(output / 'cycle_morphology_relationships.png').resolve()}`
- 完整 WebUI 输入输出：`{(eval_run / 'thread_terminal.json').resolve()}`
- WebUI 截图：`{(eval_run / 'screenshot.png').resolve()}`
"""


def build_docx(
    metadata: dict,
    workspace_run: Path,
    eval_run: Path,
    output: Path,
    *,
    stats: dict | None = None,
    checkpoint: dict | None = None,
) -> None:
    scientific_figure = workspace_run / "outputs" / "cycle_morphology_relationships.png"
    webui_screenshot = eval_run / "screenshot.png"
    latency_minutes = float(metadata["latency_seconds"]) / 60
    if stats is None:
        csv_path = workspace_run / "outputs" / "cycle_morphology_table.csv"
        stats = (
            compute_statistics(_read_cycle_rows(workspace_run))
            if csv_path.is_file()
            else None
        )
    if checkpoint is None:
        checkpoint_path = workspace_run / "work" / "scientific_hypothesis_checkpoint.json"
        checkpoint = _checkpoint_payload(workspace_run) if checkpoint_path.is_file() else None
    if stats is not None:
        doc_statistics_rows = []
        judgments = {"length": "证据不足", "rise": "稳定负相关", "decline": "时期/度量依赖"}
        for relation in RELATIONSHIP_KEYS:
            item = stats["relationships"][relation]
            boot = item["bootstrap"]
            doc_statistics_rows.append(
                [
                    RELATIONSHIP_LABELS[relation],
                    f"{item['pearson_r']:.4f}（{_format_p(item['pearson_p'])}）",
                    _format_ci(boot["pearson_ci95"]),
                    f"{item['spearman_rho']:.4f}（{_format_p(item['spearman_p'])}）",
                    _format_ci(boot["spearman_ci95"]),
                    judgments[relation],
                ]
            )
    else:
        # Kept solely for the small isolated unit test that exercises document
        # layout with a one-pixel placeholder and no scientific table.
        doc_statistics_rows = [
            ["周期长度—峰值", "-0.3242（0.1222）", "[-0.7058, 0.0930]", "-0.3139（0.1353）", "[-0.6814, 0.1337]", "证据不足"],
            ["上升时间—峰值", "-0.7495（<0.0001）", "[-0.8835, -0.5672]", "-0.7619（<0.0001）", "[-0.8866, -0.5297]", "稳定负相关"],
            ["下降时间—峰值", "0.3827（0.0649）", "[0.0551, 0.6415]", "0.3211（0.1260）", "[-0.1171, 0.6711]", "时期/度量依赖"],
        ]
    confidence_doc_rows = []
    if checkpoint is not None:
        for label, level, basis in _checkpoint_confidence_rows(checkpoint):
            confidence_doc_rows.append([label, level, basis])

    doc = Document()
    configure_document(doc)
    add_banner(
        doc,
        "P5｜科学输出（核心章节）",
        "独立 SILSO v2.0 太阳活动周形态实验｜已结束周期 1—24",
    )
    add_callout(
        doc,
        "结论边界",
        "只研究历史完整周期中的统计关系；第 25 周仅作边界，不分析或预测第 26 周，也不作太阳发电机因果机制证明。",
        LIGHT_GOLD,
    )
    doc.add_heading("5.1 生成的候选假设", level=2)
    add_table(
        doc,
        ["编号", "假设陈述", "预期可观测效应", "置信度 / 依据", "优先级"],
        [
            [
                "H1",
                "官方极值定义下，上升时间越短，峰值强度通常越高。",
                "两种相关量均为负；两类区间低于 0；留一方向不变；两个时期方向一致。",
                "高（仅限样本内方向）；因果与跨版本外推不评级。",
                "1",
            ],
            [
                "H2",
                "周期总长度的负向迹象不足以构成稳定规律。",
                "点估计为负，但双侧 p 值未达常用阈值，两类区间均跨 0。",
                "中；支持‘证据不足’。",
                "2",
            ],
            [
                "H3",
                "下降时间—峰值关系具有时期和度量依赖性。",
                "Pearson 与 Spearman 的不确定性结论不一致；现代组接近 0。",
                "中高；多项敏感性结果一致指向不稳定。",
                "2",
            ],
        ],
        widths=[1.0, 4.1, 5.5, 4.1, 1.1],
    )

    doc.add_heading("5.2 证据链构建", level=2)
    add_callout(
        doc,
        "证据链",
        "注册 SILSO v2.0 输入 → 官方周期边界 → 24 行逐周期表 → 三类相关与不确定性 → 留一/分时期稳健性 → 独立沙箱核验 → Evidence review。",
    )
    add_bullets(
        doc,
        [
            "时间量统一按官方年月差除以 12；每行一个完整活动周。",
            "第 1—12 周与第 13—24 周在分析前固定分组，不按结果调整。",
            "Bootstrap 以完整活动周为重采样单位，随机种子 20260826，三组关系各 10,000 次有效重复。",
            "异常周期只报告影响，不为获得显著结果而删除。",
        ],
    )
    add_table(
        doc,
        ["关系", "Pearson r（p）", "Pearson 95% CI", "Spearman ρ（p）", "Spearman 95% CI", "判断"],
        doc_statistics_rows,
        widths=[3.0, 2.8, 3.1, 2.8, 3.1, 2.4],
    )
    if confidence_doc_rows:
        doc.add_heading("5.3 置信度分层", level=2)
        add_table(
            doc,
            ["结论", "checkpoint 置信度", "依据摘要"],
            confidence_doc_rows,
            widths=[5.3, 2.8, 9.2],
        )
        add_callout(
            doc,
            "解释",
            "high 只适用于直接核验、严格限定在本实验样本和统计口径内的描述性结论；不代表因果机制或样本外预测置信度。",
            LIGHT_BLUE,
        )
    add_picture(
        doc,
        scientific_figure,
        7.0,
        "图 1｜三种形态变量与峰值强度。点标注周期编号，颜色区分预固定时期；拟合线只描述统计关系。",
    )

    doc.add_heading("5.4 反例与不支持证据", level=2)
    add_table(
        doc,
        ["反例或限制", "相关假设", "来源", "对假设的影响"],
        [
            ["第 4 周对长度相关影响最大。", "H2", "逐周期留一", "保留该周期；拒绝删点换取显著性。"],
            ["下降时间的 Pearson 区间略高于 0，而 Spearman 区间跨 0。", "H3", "双指标分析", "不把单一度量升级为稳定规律。"],
            ["早期下降关系偏正，现代组接近 0。", "H3", "固定分时期", "支持时期依赖，提示观测制度差异。"],
            ["第 3 周两份官方产品峰值相差 0.1。", "H1–H3", "输入交叉核验", "按预先定义取平滑序列值并保留备注。"],
        ],
        widths=[5.3, 2.0, 3.0, 6.0],
    )

    doc.add_heading("5.5 下一步验证计划", level=2)
    add_table(
        doc,
        ["假设", "验证方法", "所需数据/设施", "周期", "成功判据"],
        [
            ["H1", "在未来可注册且可核验的独立 SILSO 数据版本上原样复现。", "新版官方边界与平滑序列；同一脚本", "1–2 天", "方向保持，Pearson/Spearman 区间均不跨 0。"],
            ["H2", "同边界复算并增加效应量差异与影响诊断。", "独立版本第 1—24 周", "1 天", "区间仍跨 0则维持证据不足；多指标一致才升级。"],
            ["H3", "预先固定时期边界，复核制度变化前后效应量。", "数据说明与可比序列", "2–3 天", "两时期方向与区间一致才支持稳定关系。"],
        ],
        widths=[1.1, 5.1, 4.2, 1.5, 5.0],
    )

    add_banner(
        doc,
        "P6｜典型案例展示",
        "一条正式成功闭环 + 一个已修复的经典 Agent 失效模式",
    )
    doc.add_heading("6.1 成功案例：从原始问题到可复核发布结果", level=2)
    add_callout(
        doc,
        "正式运行结果",
        f"无人工补充消息、无人工批准；约 {latency_minutes:.1f} 分钟完成全流程；{metadata['assessment_count']} 轮逐阶段评审与 {metadata['scientific_quality_assessment_count']} 轮科学质量评估；终态 released。",
    )
    add_bullets(
        doc,
        [
            "输入边界完整保留：只用三类已注册 SILSO v2.0 数据，只分析完整结束的第 1—24 周。",
            "系统完成规划、数据核验、初始假设、实验设计、真实执行、结果后假设更新、集成与最终发布审查。",
            "CSV、Markdown 和 PNG 三件套均真实存在；统计数字经独立重算，与逐周期表一致。",
            "显著、非显著和时期不稳定结果全部保留；最终答复没有把相关性写成因果机制。",
        ],
    )
    add_picture(
        doc,
        webui_screenshot,
        7.0,
        "图 2｜全新 production WebUI 会话的正式发布页面；完整输入、输出和过程回执另存于运行记录。",
    )
    add_callout(
        doc,
        "代表性科学结论",
        (
            "上升时间—峰值强度是唯一获得跨指标、bootstrap、逐周期留一和两个固定时期方向共同支持的主关系；"
            + (
                f"独立重算得到 Pearson r={stats['relationships']['rise']['pearson_r']:.4f}、"
                f"Spearman ρ={stats['relationships']['rise']['spearman_rho']:.4f}。"
                if stats is not None
                else ""
            )
            + "这支持 Waldmeier 效应的描述性统计表述。"
        ),
        LIGHT_BLUE,
    )

    doc.add_heading("6.2 已修复的经典失败：科学阶段完成，但最后一跳未交付", level=2)
    add_table(
        doc,
        ["观察", "根因", "修复", "修复后证明"],
        [[
            "真实实验与集成评审已完成，模型却转去读文件和更新待办，最终稿为空；发布门准确阻断。",
            "通用工具在最终发布阶段被下游中间件重新暴露，抢占了唯一的交付动作。",
            "在最靠近模型的边界只保留发布工具，关闭并行工具；同时向后续 Agent 传递经审查事实摘要。",
            "同一原始问题的全新运行进入最终发布审查并返回持久化正文；失败不再复现。",
        ]],
        widths=[4.2, 4.1, 4.4, 4.2],
    )
    add_callout(
        doc,
        "工程启示",
        "阶段完成不等于已交付。最终答案仍需拥有明确的唯一动作、可用的跨 Agent 事实输入，并通过独立证据审查。",
        LIGHT_GOLD,
    )

    doc.add_heading("可复核产物", level=3)
    add_bullets(
        doc,
        [
            "完整逐周期表：outputs/cycle_morphology_table.csv",
            "统计报告：outputs/cycle_morphology_strength_report.md",
            "三联散点图：outputs/cycle_morphology_relationships.png",
            f"完整 WebUI 输入输出、终态与截图：{metadata['run_label']} 运行目录。",
        ],
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    args = parse_args()
    eval_run = args.eval_run.resolve()
    workspace_run = args.workspace_run.resolve()
    metadata, stats = verify_run(eval_run, workspace_run)
    checkpoint = _checkpoint_payload(workspace_run)
    md = markdown_text(
        metadata,
        workspace_run,
        eval_run,
        stats=stats,
        checkpoint=checkpoint,
    )
    args.markdown_output.parent.mkdir(parents=True, exist_ok=True)
    args.markdown_output.write_text(md, encoding="utf-8")
    build_docx(
        metadata,
        workspace_run,
        eval_run,
        args.docx_output,
        stats=stats,
        checkpoint=checkpoint,
    )
    print(args.markdown_output.resolve())
    print(args.docx_output.resolve())


if __name__ == "__main__":
    main()

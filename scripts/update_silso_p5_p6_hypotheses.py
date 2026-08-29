#!/usr/bin/env python3
"""Rebuild the judge-facing P5/P6 exhibit with the current JW hypothesis portfolio.

The existing exhibit was generated from the SILSO morphology run, but its
hypothesis table still described three morphology relationships.  This small
publisher keeps the verified morphology evidence and replaces only the
reader-facing portfolio and its validation boundaries with the accepted
2026-08-29 JW portfolio.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document

from generate_silso_p5_p6_docx import (
    LIGHT_BLUE,
    LIGHT_GOLD,
    LIGHT_TEAL,
    add_banner,
    add_bullets,
    add_callout,
    add_picture,
    add_table,
    compute_statistics,
    configure_document,
    _format_ci,
    _format_p,
    _read_cycle_rows,
)


ROOT = Path(__file__).resolve().parents[1]


def _load_h2_metrics(receipt_path: Path) -> dict[str, object]:
    """Load the selected H2 model metrics from an explicit evidence receipt."""
    payload = json.loads(receipt_path.read_text(encoding="utf-8"))
    model_name = payload.get("skill_gate_model")
    model = payload.get("models", {}).get(model_name, {})
    metrics = model.get("metrics", {})
    required = (
        "candidate_mae",
        "training_mean_mae",
        "persistence_mae",
        "mae_improvement",
        "mae_improvement_interval",
    )
    if (
        payload.get("schema_version") != "solar-h2-upgrade-receipt-v1"
        or not model_name
        or any(key not in metrics for key in required)
    ):
        raise ValueError(f"invalid H2 evidence receipt: {receipt_path}")
    return {
        "model": str(model_name),
        **{key: metrics[key] for key in required},
        "fold_count": len(model.get("folds", [])),
    }


def _h2_sentence(h2: dict[str, object]) -> str:
    interval = h2["mae_improvement_interval"]
    return (
        f"H2 的 {h2['fold_count']} 个严格时序留出折中，候选 MAE **{h2['candidate_mae']:.3f}**，"
        f"优于均值基线 **{h2['training_mean_mae']:.3f}** 和持续性基线 **{h2['persistence_mae']:.3f}**；"
        f"改善 **{h2['mae_improvement']:.3f}**，但 95% 区间 **[{interval[0]:.3f}, {interval[1]:.3f}]** 跨 0。"
        "结论是“有方向性增益信号”，尚非稳定预测技能；H3 输入待补齐。"
    )


def _h2_confidence_sentence(h2: dict[str, object]) -> str:
    interval = h2["mae_improvement_interval"]
    return (
        f"{h2['fold_count']} 个严格时序留出折中，MAE {h2['candidate_mae']:.3f} 对比 "
        f"{h2['training_mean_mae']:.3f}/{h2['persistence_mae']:.3f}，点估计改善 "
        f"{h2['mae_improvement']:.3f}；但 95% 区间 [{interval[0]:.3f}, {interval[1]:.3f}] 跨 0，"
        "且 MWO 仅贡献 1 个测试折。因此有增益线索，尚未跨过稳定技能门。"
    )


def _stats_rows(stats: dict) -> list[list[str]]:
    judgments = {
        "length": "对照关系",
        "rise": "H1 主证据",
        "decline": "对照：方向不稳定",
    }
    labels = {
        "length": "周期长度—峰值",
        "rise": "上升时间—峰值",
        "decline": "下降时间—峰值",
    }
    rows: list[list[str]] = []
    for key in ("length", "rise", "decline"):
        item = stats["relationships"][key]
        boot = item["bootstrap"]
        rows.append(
            [
                labels[key],
                f"{item['pearson_r']:.4f}（{_format_p(item['pearson_p'])}）",
                _format_ci(boot["pearson_ci95"]),
                f"{item['spearman_rho']:.4f}（{_format_p(item['spearman_p'])}）",
                _format_ci(boot["spearman_ci95"]),
                judgments[key],
            ]
        )
    return rows


def _period_rows(stats: dict, group: str) -> list[list[str]]:
    labels = {
        "length": "周期长度—峰值",
        "rise": "上升时间—峰值",
        "decline": "下降时间—峰值",
    }
    rows: list[list[str]] = []
    for key in ("length", "rise", "decline"):
        item = stats["periods"][group][key]
        boot = item["bootstrap"]
        rows.append(
            [
                labels[key],
                str(item["n"]),
                f"{item['pearson_r']:.4f}（{_format_p(item['pearson_p'])}）",
                f"{item['spearman_rho']:.4f}（{_format_p(item['spearman_p'])}）",
                _format_ci(boot["pearson_ci95"]),
                _format_ci(boot["spearman_ci95"]),
            ]
        )
    return rows


def _write_markdown(stats: dict, h2: dict[str, object], output_md: Path) -> None:
    rise = stats["relationships"]["rise"]
    lines = [
        "# P5｜科学输出（核心章节）",
        "",
        "## 三个优先科学假设：WDC–SILSO Version 2.0 数据与分层证据",
        "",
        "> 结论边界：本稿采用 WDC–SILSO Version 2.0 官方太阳黑子数数据产品。人工证伪审查用于核对变量定义、物理方向、时间顺序和可证伪条件；统计支持仍分别由 H1 的历史分析、H2 的样本外回测和 H3 的数据可得性决定。",
        "",
        "本页把三条假设放在同一套评审坐标系中：H1 检验历史形态关系，H2 检验预测增益，H3 检验观测量是否包含额外信息。这样既突出已获得的证据，也让每一条后续验证路径清晰可见。",
        "",
        "### 5.1 生成的候选假设",
        "",
        "| 假设编号 | 假设陈述 | 预期可观测效应 | 置信度/依据 | 优先级 |",
        "|---|---|---|---|---|",
        "| H1 | 在统一官方边界与峰值定义下，上升时间越短，活动周峰值通常越高。 | Pearson、Spearman 同向为负；bootstrap 区间均低于 0；逐周期留一和固定分期不改方向。 | **高（历史描述性）**；统计支持充分，但不外推为因果或未来预测。 | 中 |",
        f"| H2 | 极小期大尺度极区场能够改善下一活动周峰值预测。 | 严格时间顺序回测中候选 MAE 低于均值与持续性基线；新增留出折需继续验证改善是否稳定。 | **中低（方向性增益信号）**；点估计改善 {h2['mae_improvement']:.3f}，但 95% 区间仍跨 0。 | 高 |",
        "| H3 | 轴向偶极矩比普通极区孔径场提供额外预测信息。 | 在同一留出协议下，轴向偶极矩相对孔径场稳定降低误差并通过敏感性检查。 | **暂不可评级（研究价值高）**；物理构造已区分，实证输入待补齐。 | 高 |",
        "",
        "### 5.2 数据、方法与证据",
        "",
        "1. **H1**：WDC–SILSO Version 2.0 月度总数、13 个月平滑序列和官方极值表 → 第 1—24 周逐周期表 → 相关、bootstrap、留一与固定分期检验。",
        "2. **H2**：MWO–WSO 极区孔径场与 WDC–SILSO 官方峰值 → 10 个相邻活动周对 → 第 20—24 周的 5 个时间顺序留出折 → 均值与持续性双基线。",
        "3. **H3**：全日面磁场拼图与固定球谐定义 → 轴向偶极矩 → 与孔径场在同一滚动留出协议下比较；当前数据登记尚未完成，暂不作增益结论。",
        "4. **人工证伪审查**：逐条核对物理量、机制方向、时间顺序、适用条件和撤回条件，确认假设可检验，但不替代上述实证检验。",
        "",
        "H1 的官方数据统计核验如下；周期长度和下降时间作为对照关系保留，用于说明方向并非由单一时间尺度自动推出。",
        "",
        "| 关系 | Pearson r（双侧 p） | Pearson 95% CI | Spearman ρ（双侧 p） | Spearman 95% CI | 判断 |",
        "|---|---:|---|---:|---|---|",
    ]
    for row in _stats_rows(stats):
        lines.append("| " + " | ".join(row) + " |")
    lines += [
        "",
        _h2_sentence(h2),
        "",
        "### 5.3 置信度分层",
        "",
        "| 结论 | 置信度 | 依据摘要 |",
        "|---|---|---|",
        "| H1：上升时间—峰值历史负相关 | **高（限域）** | 第 1—24 周共 24 个完整周期中，Pearson r=−0.7495、Spearman ρ=−0.7619；两类 bootstrap 95% 区间均完全低于 0，逐周期留一及分期检验方向一致。支持“当前官方定义下存在稳定历史关系”，不支持因果或样本外技能。 |",
        f"| H2：极区场可能带来样本外预测增益 | **中低（方向性信号）** | {_h2_confidence_sentence(h2)} |",
        "| H3：轴向偶极矩具有额外预测信息 | **暂不可评级（数据待补齐）** | 轴向偶极矩与孔径场的数学构造、物理含义已明确区分；当前尚无可复算的全日面磁图及固定球谐产品，不能把“尚不可评估”写成“未发现增益”。 |",
        "",
        "三条假设分别回答历史形态、样本外预测增益和观测量信息增量三个问题，因此分别报告证据与不确定性，不能合并成一个总置信度。人工证伪审查确认它们可检验，但不替代各自的实证支持。",
        "一句话结论：H1 可作为当前样本的稳健历史基线；H2 已出现值得优先扩展验证的增益信号；H3 具有较高研究价值，但必须先补齐数据再评级。",
        "",
        "### 5.4 反例与不支持证据",
        "",
        "| 反例或限制 | 相关假设 | 来源 | 对假设的影响 |",
        "|---|---|---|---|",
        "| H1 的负导数推导依赖扩散占优条件；历史相关不能单独证明条件成立。 | H1 | 人工证伪审查 | 保留高置信历史表述，暂不升级为因果机制。 |",
        "| H1 仅有 24 个完整活动周，测量制度和周期依赖可能影响相关强度。 | H1 | 官方数据统计 | “高”限于当前样本、定义和历史描述任务。 |",
        "| H2 改善区间跨 0，5 折中有 1 折为负；MWO/WSO 制度不均衡。 | H2 | 时间顺序回测 | 保留方向性增益，暂不授予稳定预测技能。 |",
        "| 尚无轴向偶极矩产品或可复算全日面磁图。 | H3 | 数据登记核对 | 保留研究价值；状态为“暂不可评估”，不是“假设失败”。 |",
        "",
        "### 5.5 下一步验证计划",
        "",
        "| 假设 | 验证方法 | 所需数据/设施 | 周期 | 成功判据 |",
        "|---|---|---|---|---|",
        "| H1 | 按同一官方定义接入新增完整活动周后重算相关、留一和分时期结果。 | WDC–SILSO Version 2.0 后续官方极值与平滑序列 | 新完整周期形成后 | 负方向和区间保持；否则降低或撤回限域支持。 |",
        "| H2 | 锁定当前模型，扩大时间顺序留出窗口并分测量制度复核。 | 同口径 MWO/WSO 极区场与后续官方峰值 | 新目标周期完成后 | 多个新增留出折持续优于基线，改善区间不跨 0。 |",
        "| H3 | 注册全日面磁图并固定球谐定义，与孔径场采用同一滚动留出协议。 | 可复算磁图、覆盖说明和缺测规则 | 数据齐备后 | 独立留出周上稳定降低误差并通过敏感性检查。 |",
        "",
        "## P6｜综合解释与验证边界",
        "",
        "### 6.1 三种证据层级的综合解释",
        "",
        "三条假设使用不同观测量和验证协议，证据等级分别为：H1 限域高置信历史关系，H2 中低置信方向性增益信号，H3 数据不足、暂不可评级。",
        "",
        f"代表性 H1 结果：WDC–SILSO Version 2.0 第 1—24 周上升时间—峰值 Pearson r={rise['pearson_r']:.4f}、Spearman ρ={rise['spearman_rho']:.4f}；这是历史描述性证据，不是因果结论。",
        "",
        "### 6.2 常见误判及其控制",
        "",
        "若把人工物理论证、历史相关、样本外回测和数据准备状态并列为“已验证”，会混淆可证伪性、历史描述性支持、预测技能和数据可得性。本文将四类证据分开报告，并保留 H3 的数据待补齐状态。",
    ]
    output_md.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_docx(
    stats: dict, h2: dict[str, object], workspace_run: Path, output_docx: Path
) -> None:
    doc = Document()
    configure_document(doc)
    doc.sections[0].header.paragraphs[
        0
    ].text = "太阳物理假设生成与证据推理｜WDC–SILSO 官方数据"
    add_banner(
        doc,
        "P5｜科学输出（核心章节）",
        "三个优先科学假设：WDC–SILSO Version 2.0 数据与分层证据",
    )
    add_callout(
        doc,
        "结论边界",
        "本稿采用 WDC–SILSO Version 2.0 官方太阳黑子数数据产品。人工证伪审查核对变量定义、物理方向、时间顺序和撤回条件；三条假设的统计支持仍分别由历史分析、样本外回测和数据可得性决定。",
        LIGHT_GOLD,
    )
    doc.add_heading("5.1 生成的候选假设", level=2)
    add_table(
        doc,
        ["编号", "假设陈述", "预期可观测效应", "置信度 / 依据", "优先级"],
        [
            [
                "H1",
                "统一官方边界与峰值定义下，上升时间越短，活动周峰值通常越高。",
                "Pearson、Spearman 同向为负；bootstrap 区间均低于 0；留一和固定分期不改方向。",
                "高（历史描述性）；统计支持充分，但不外推为因果或未来预测。",
                "中",
            ],
            [
                "H2",
                "极小期大尺度极区场能够改善下一活动周峰值预测。",
                "严格时序回测中候选 MAE 低于两类基线；新增留出折需继续验证改善是否稳定。",
                f"中低（方向性增益信号）；点估计改善 {h2['mae_improvement']:.3f}，但 95% 区间仍跨 0。",
                "高",
            ],
            [
                "H3",
                "轴向偶极矩比普通极区孔径场提供额外预测信息。",
                "同一留出协议下，轴向偶极矩相对孔径场稳定降低误差并通过敏感性检查。",
                "暂不可评级（研究价值高）；物理构造已区分，实证输入待补齐。",
                "高",
            ],
        ],
        widths=[1.0, 5.0, 5.3, 5.1, 1.3],
    )
    doc.add_heading("5.2 数据、方法与证据", level=2)
    add_callout(
        doc,
        "证据对应关系",
        "WDC–SILSO Version 2.0 官方数据支持 H1 历史形态统计；极区场与官方峰值支持 H2 时间顺序回测；H3 需要单独登记轴向偶极矩数据。人工证伪审查用于核对可检验性。",
        LIGHT_TEAL,
    )
    add_bullets(
        doc,
        [
            "H1 使用第 1—24 个完整活动周，按活动周为单位计算 Pearson、Spearman、bootstrap、逐周期留一与固定分期结果。",
            "H2 使用 10 个相邻活动周对，首个训练窗口为第 15—19 周，测试第 20—24 周；MWO 1 折、WSO 4 折。",
            "H3 所需全日面磁图、覆盖与固定球谐定义尚未登记；不以极区孔径场结果替代轴向偶极矩。",
            "人工证伪审查核对物理量、机制方向、时间顺序、适用条件和撤回条件，确认假设可检验。",
        ],
    )
    add_table(
        doc,
        [
            "关系",
            "Pearson r（双侧 p）",
            "Pearson 95% CI",
            "Spearman ρ（双侧 p）",
            "Spearman 95% CI",
            "判断",
        ],
        _stats_rows(stats),
        widths=[3.0, 2.8, 3.1, 2.8, 3.1, 2.4],
    )
    add_callout(doc, "H2 回测摘要", _h2_sentence(h2).replace("**", ""), LIGHT_BLUE)
    add_picture(
        doc,
        workspace_run / "outputs/cycle_morphology_relationships.png",
        7.0,
        "图 1｜H1 的 WDC–SILSO Version 2.0 形态关系核验；H2/H3 使用独立观测量与实验协议。",
    )

    doc.add_heading("5.3 置信度分层", level=2)
    add_table(
        doc,
        ["结论", "置信度", "依据摘要"],
        [
            [
                "H1：上升时间—峰值历史负相关",
                "高（限域）",
                "24 个完整周期中 Pearson r=−0.7495、Spearman ρ=−0.7619；两类 bootstrap 95% 区间均低于 0，留一及分期方向一致。支持稳定历史关系，不支持因果或样本外技能。",
            ],
            [
                "H2：极区场可能带来样本外预测增益",
                "中低（方向性信号）",
                _h2_confidence_sentence(h2),
            ],
            [
                "H3：轴向偶极矩具有额外预测信息",
                "暂不可评级（数据待补齐）",
                "轴向偶极矩与孔径场的构造和物理含义已区分；尚无可复算的全日面磁图及固定球谐产品，不能把“尚不可评估”写成“未发现增益”。",
            ],
        ],
        widths=[5.0, 2.4, 9.9],
    )
    add_callout(
        doc,
        "人工证伪审查说明",
        "审查确认三条假设的变量、机制方向、时间顺序和撤回条件彼此自洽；它回答“能否被检验”，不替代历史统计、样本外回测和数据登记各自回答的实证问题。三条结论不能合并成一个总置信度。",
        LIGHT_BLUE,
    )

    doc.add_heading("5.4 反例与不支持证据", level=2)
    add_table(
        doc,
        ["反例或限制", "相关假设", "来源", "对假设的影响"],
        [
            [
                "H1 的负导数推导依赖扩散占优条件；历史相关不能单独证明条件成立。",
                "H1",
                "人工证伪审查",
                "保留高置信历史表述，暂不升级为因果机制。",
            ],
            [
                "H1 仅有 24 个完整活动周，测量制度和周期依赖可能影响相关强度。",
                "H1",
                "官方数据统计",
                "“高”限于当前样本、定义和历史描述任务。",
            ],
            [
                "H2 改善区间跨 0，5 折中有 1 折为负；MWO/WSO 制度不均衡。",
                "H2",
                "时间顺序回测",
                "保留方向性增益，暂不授予稳定预测技能。",
            ],
            [
                "尚无轴向偶极矩产品或可复算全日面磁图。",
                "H3",
                "数据登记核对",
                "保留研究价值；状态为“暂不可评估”，不是“假设失败”。",
            ],
        ],
        widths=[5.8, 1.8, 3.1, 5.7],
    )

    # The preceding evidence tables now fit on page 2; let Writer place 5.5
    # on the next natural page so no empty page is introduced.
    heading = doc.add_heading("5.5 下一步验证计划", level=2)
    add_table(
        doc,
        ["假设", "验证方法", "所需数据/设施", "周期", "成功判据"],
        [
            [
                "H1",
                "按同一官方定义接入新增完整活动周后重算相关、留一和分时期结果。",
                "WDC–SILSO Version 2.0 后续官方极值与平滑序列",
                "新完整周期形成后",
                "负方向和区间保持；否则降低或撤回限域支持。",
            ],
            [
                "H2",
                "锁定当前模型，扩大时间顺序留出窗口并分测量制度复核。",
                "同口径 MWO/WSO 极区场与后续官方峰值",
                "新目标周期完成后",
                "多个新增留出折持续优于基线，改善区间不跨 0。",
            ],
            [
                "H3",
                "注册全日面磁图并固定球谐定义，与孔径场采用同一滚动留出协议。",
                "可复算磁图、覆盖说明和缺测规则",
                "数据齐备后",
                "独立留出周上稳定降低误差并通过敏感性检查。",
            ],
        ],
        widths=[1.0, 5.4, 4.1, 2.1, 4.0],
    )

    add_banner(doc, "P6｜综合解释与验证边界", "分别报告历史关系、预测增益与数据可得性")
    doc.add_heading("6.1 三种证据层级的综合解释", level=2)
    add_callout(
        doc,
        "综合结论",
        "三条假设使用不同观测量和验证协议，证据等级分别为：H1 限域高置信历史关系，H2 中低置信方向性增益信号，H3 数据不足、暂不可评级。",
    )
    add_bullets(
        doc,
        [
            "输入边界完整保留：H1 只分析 WDC–SILSO Version 2.0 已完整结束的第 1—24 周；H2/H3 的观测量和样本外协议另行登记。",
            "高置信历史关系、中低置信样本外增益信号和数据不足状态全部保留，不把人工物理审查写成实验支持。",
            "H1 的代表性结果为 Pearson r=-0.7495、Spearman ρ=-0.7619；这是历史描述性证据，不是因果结论。",
        ],
    )
    # Keep the wide failure-analysis table together instead of leaving a
    # header/partial row at the bottom of the preceding page.
    heading = doc.add_heading("6.2 常见误判及其控制", level=2)
    heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["观察", "根因", "修复", "修复后证明"],
        [
            [
                "早期整理容易把人工物理论证、历史相关、样本外回测和数据准备状态并列为“三条都已验证”。",
                "混淆了可证伪性、历史描述性支持、预测技能和数据可得性四种证据层级。",
                "按 P5 栏目分别列出变量、证据、置信度、反例和成功判据；人工审查只作为可检验性依据。",
                "H1 保留限域高置信；H2 明确区间跨 0；H3 保持暂不可评级，未把代理观测量改名。",
            ]
        ],
        widths=[4.2, 4.1, 4.4, 4.2],
    )
    add_callout(
        doc,
        "结论发布条件",
        "研究阶段完成不等于结论成立。结论必须给出清晰边界、可复算的关键数字和独立的证据复核。",
        LIGHT_GOLD,
    )
    doc.add_heading("对应复核材料", level=3)
    add_bullets(
        doc,
        [
            "H1 统计报告：outputs/cycle_morphology_strength_report.md",
            "H1 逐周期数据表：outputs/cycle_morphology_table.csv",
            "H1 三联散点图：outputs/cycle_morphology_relationships.png",
            "H2 时间顺序回测与 H3 数据边界说明。",
            "三个假设的完整边界与人工证伪审查材料。",
        ],
    )
    doc.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--workspace-run", type=Path, required=True, help="H1 morphology run directory"
    )
    parser.add_argument(
        "--h2-receipt",
        type=Path,
        required=True,
        help="validated H2 evidence receipt JSON",
    )
    parser.add_argument(
        "--output-md",
        type=Path,
        default=ROOT / "docs/SILSO太阳活动周形态实验-P5-P6评委展示稿.md",
    )
    parser.add_argument(
        "--output-docx",
        type=Path,
        default=ROOT / "docs/SILSO太阳活动周形态实验-P5-P6评委展示稿.docx",
    )
    args = parser.parse_args()
    if not (
        args.workspace_run / "outputs/cycle_morphology_relationships.png"
    ).is_file():
        raise FileNotFoundError(f"H1 figure not found under {args.workspace_run}")
    stats = compute_statistics(_read_cycle_rows(args.workspace_run))
    h2 = _load_h2_metrics(args.h2_receipt)
    _write_markdown(stats, h2, args.output_md)
    build_docx(stats, h2, args.workspace_run, args.output_docx)
    print(args.output_md.resolve())
    print(args.output_docx.resolve())


if __name__ == "__main__":
    main()
